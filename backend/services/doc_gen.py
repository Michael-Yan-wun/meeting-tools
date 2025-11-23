from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders (not directly supported by python-docx, simple pass for now)
    For a simple implementation, we rely on default table styles.
    """
    pass

def generate_meeting_minutes(data):
    """
    生成會議記錄 Word 檔案 (依照使用者指定格式)
    data: 包含 filename, created_at, participants, key_points, next_steps, summary, meeting_topics
    回傳: BytesIO 物件
    """
    doc = Document()
    
    # 設定中文字型 (雖然 python-docx 對中文字型支援有限，但盡量設定)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # --- 標題 ---
    heading_text = f"{data.get('filename', '會議記錄')}".replace(".mp3", "").replace(".wav", "") + " 會議記錄整理"
    heading = doc.add_heading(heading_text, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 一、會議主題 ---
    doc.add_heading("一、會議主題", level=1)
    topics = data.get('meeting_topics', [])
    if topics:
        for topic in topics:
            p = doc.add_paragraph(str(topic), style='List Bullet')
    else:
        doc.add_paragraph("(無特定主題)")

    # --- 二、參與人物 (表格) ---
    doc.add_heading("二、參與人物", level=1)
    participants = data.get('participants', [])
    
    # 建立表格：標題列 + 內容
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # 使用網格樣式
    
    # 設定表頭
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '姓名 / 角色'
    hdr_cells[1].text = '身份或貢獻說明'
    
    if participants:
        for p_data in participants:
            row_cells = table.add_row().cells
            # 處理舊格式 (如果是字串列表) 或新格式 (字典列表)
            if isinstance(p_data, dict):
                row_cells[0].text = str(p_data.get('name', ''))
                row_cells[1].text = str(p_data.get('role', ''))
            else:
                row_cells[0].text = str(p_data)
                row_cells[1].text = ""
    else:
        row_cells = table.add_row().cells
        row_cells[0].text = "(無資料)"
        row_cells[1].text = ""

    # --- 三、重點內容摘要 ---
    doc.add_heading("三、重點內容摘要", level=1)
    key_points = data.get('key_points', [])
    
    if key_points:
        for idx, kp in enumerate(key_points, 1):
            if isinstance(kp, dict):
                # 新格式: {"title": "...", "content": "..."}
                title = kp.get('title', '')
                content = kp.get('content', '')
                # 標題段落
                p_title = doc.add_paragraph(f"{idx}. {title}", style='List Number')
                # 內容段落 (縮排)
                if content:
                    p_content = doc.add_paragraph(content)
                    p_content.paragraph_format.left_indent = Inches(0.25)
            else:
                # 舊格式相容 (純字串)
                doc.add_paragraph(str(kp), style='List Bullet')
    else:
        doc.add_paragraph("(無重點摘要)")

    # --- 四、會議結論與行動項目 (表格) ---
    doc.add_heading("四、會議結論與行動項目", level=1)
    next_steps = data.get('next_steps', [])
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '行動項目'
    hdr_cells2[1].text = '負責人 / 協調人'
    
    if next_steps:
        for step in next_steps:
            row_cells = table2.add_row().cells
            if isinstance(step, dict):
                row_cells[0].text = str(step.get('action', ''))
                row_cells[1].text = str(step.get('owner', ''))
            else:
                row_cells[0].text = str(step)
                row_cells[1].text = ""
    else:
        row_cells = table2.add_row().cells
        row_cells[0].text = "(無待辦事項)"
        row_cells[1].text = ""

    # --- 五、總結摘要 ---
    doc.add_heading("五、總結摘要", level=1)
    summary = data.get('summary', '')
    if summary:
        doc.add_paragraph(summary)
    else:
        doc.add_paragraph("(無總結)")

    # 移除逐字稿區段 (根據使用者要求)

    # 儲存到 BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
