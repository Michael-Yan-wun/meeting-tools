from docx import Document
import sys

def inspect_docx(filename):
    try:
        doc = Document(filename)
        print(f"--- 分析檔案: {filename} ---\n")
        
        print("[段落結構]")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                print(f"段落 {i} (Style: {p.style.name}): {p.text[:50]}...")
        
        print("\n[表格結構]")
        for i, table in enumerate(doc.tables):
            print(f"表格 {i}: {len(table.rows)} 列 x {len(table.columns)} 欄")
            # 印出前幾列內容看看
            for r_idx, row in enumerate(table.rows[:2]):
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                print(f"  Row {r_idx}: {row_text[:50]}...")

    except Exception as e:
        print(f"讀取失敗: {e}")

if __name__ == "__main__":
    inspect_docx("中正北路24號_會議記錄整理.docx")

